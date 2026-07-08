"""Ingestion: turn a contract file into text and coarse clause sections.

The demo ships plain-text contracts so the whole thing runs offline with no
parsing dependencies. Production would drop a PDF/DOCX extractor in here
(pdfplumber / python-docx); nothing downstream changes because everything works
off ``Document.text`` and ``Document.sections``.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

# A heading line looks like "3. Confidential Information": a number, then a
# short Title-Case phrase on its own line (no trailing full stop).
_HEADING = re.compile(r"^\s*(\d+)[.)]\s+([A-Z][A-Za-z0-9 &/,'-]{1,48})\s*$")


@dataclass
class Section:
    number: int
    heading: str
    body: str


@dataclass
class Document:
    path: str
    title: str
    text: str
    sections: List[Section] = field(default_factory=list)


def load(path: str | Path) -> Document:
    """Read a contract file and split it into titled sections."""
    p = Path(path)
    raw = _read_any(p)
    lines = raw.splitlines()

    title = next((ln.strip() for ln in lines if ln.strip()), p.stem)

    sections: List[Section] = []
    current = None
    buf: List[str] = []
    for ln in lines:
        m = _HEADING.match(ln)
        if m:
            if current is not None:
                current.body = "\n".join(buf).strip()
                sections.append(current)
            current = Section(number=int(m.group(1)), heading=m.group(2).strip(), body="")
            buf = []
        elif current is not None:
            buf.append(ln)
    if current is not None:
        current.body = "\n".join(buf).strip()
        sections.append(current)

    return Document(path=str(p), title=title, text=raw, sections=sections)


def _read_any(p: Path) -> str:
    suffix = p.suffix.lower()
    if suffix in (".txt", ".md", ""):
        return p.read_text(encoding="utf-8", errors="ignore")
    # Hooks for production formats. Kept optional so the demo has zero deps.
    if suffix == ".pdf":  # pragma: no cover - optional dependency
        import pdfplumber  # type: ignore

        with pdfplumber.open(str(p)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    if suffix == ".docx":  # pragma: no cover - optional dependency
        import docx  # type: ignore

        return "\n".join(par.text for par in docx.Document(str(p)).paragraphs)
    raise ValueError(f"Unsupported file type: {suffix}")
